"""Convert PRD and Dev Plan markdown files into formatted Word documents."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOCS_DIR = Path(__file__).parent
BRAND_NAVY = RGBColor(0x1E, 0x3A, 0x5F)
GRAY = RGBColor(0x55, 0x55, 0x55)


def add_heading(doc, text, level):
    """Add a styled heading."""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_NAVY
        if level == 0:
            run.font.size = Pt(28)
        elif level == 1:
            run.font.size = Pt(20)
        elif level == 2:
            run.font.size = Pt(16)
        elif level == 3:
            run.font.size = Pt(13)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def parse_inline(text):
    """Parse **bold** and *italic* and `code` markers — return list of (text, bold, italic, code)."""
    parts = []
    pattern = re.compile(r'(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)')
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            parts.append((text[pos:m.start()], False, False, False))
        token = m.group(0)
        if token.startswith('**'):
            parts.append((token[2:-2], True, False, False))
        elif token.startswith('`'):
            parts.append((token[1:-1], False, False, True))
        else:
            parts.append((token[1:-1], False, True, False))
        pos = m.end()
    if pos < len(text):
        parts.append((text[pos:], False, False, False))
    return parts


def add_rich_para(doc, text, style=None, size=11):
    p = doc.add_paragraph(style=style)
    for chunk, bold, italic, code in parse_inline(text):
        run = p.add_run(chunk)
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Consolas"
            run.font.size = Pt(10)
    return p


def add_table_from_md(doc, header_row, data_rows):
    """Create a Word table from parsed markdown header + data rows."""
    table = doc.add_table(rows=1 + len(data_rows), cols=len(header_row))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header_row):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h.strip())
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    # Header shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in hdr:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1E3A5F')
        tc_pr.append(shd)
    for r_idx, row in enumerate(data_rows, start=1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            if c_idx < len(cells):
                cells[c_idx].text = ""
                p = cells[c_idx].paragraphs[0]
                for chunk, bold, italic, code in parse_inline(val.strip()):
                    run = p.add_run(chunk)
                    run.font.size = Pt(10)
                    run.bold = bold
                    run.italic = italic


def md_to_docx(md_path, docx_path, doc_title):
    """Convert a markdown file to a styled docx."""
    text = Path(md_path).read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()
    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # Cover title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(doc_title)
    title_run.bold = True
    title_run.font.size = Pt(26)
    title_run.font.color.rgb = BRAND_NAVY

    sub = doc.add_paragraph()
    sub_run = sub.add_run("SaralPrivacy.com")
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = GRAY

    doc.add_paragraph()  # spacer

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines (but preserve paragraph breaks)
        if not stripped:
            i += 1
            continue

        # Headings
        if stripped.startswith("# "):
            add_heading(doc, stripped[2:].strip(), 0)
            i += 1
            continue
        if stripped.startswith("## "):
            add_heading(doc, stripped[3:].strip(), 1)
            i += 1
            continue
        if stripped.startswith("### "):
            add_heading(doc, stripped[4:].strip(), 2)
            i += 1
            continue
        if stripped.startswith("#### "):
            add_heading(doc, stripped[5:].strip(), 3)
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            doc.add_paragraph()
            i += 1
            continue

        # Code blocks
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            run = p.add_run("\n".join(code_lines))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            # subtle gray bg via paragraph format would need shading — skip for simplicity
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(r'^\|[\s\-:|]+\|$', lines[i+1].strip()):
            header = [c.strip() for c in stripped.strip("|").split("|")]
            i += 2  # skip separator
            data = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = [c.strip() for c in lines[i].strip().strip("|").split("|")]
                data.append(row)
                i += 1
            add_table_from_md(doc, header, data)
            continue

        # Bullets
        if stripped.startswith("- ") or stripped.startswith("* "):
            add_rich_para(doc, stripped[2:].strip(), style="List Bullet")
            i += 1
            continue

        # Numbered
        if re.match(r'^\d+\.\s', stripped):
            content = re.sub(r'^\d+\.\s', '', stripped)
            add_rich_para(doc, content, style="List Number")
            i += 1
            continue

        # Plain paragraph
        add_rich_para(doc, stripped)
        i += 1

    doc.save(docx_path)
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    md_to_docx(
        DOCS_DIR / "SaralPrivacy_PRD.md",
        DOCS_DIR / "SaralPrivacy_PRD.docx",
        "Product Requirements Document",
    )
    md_to_docx(
        DOCS_DIR / "SaralPrivacy_DevelopmentPlan.md",
        DOCS_DIR / "SaralPrivacy_DevelopmentPlan.docx",
        "Development Plan — Block-wise",
    )
    md_to_docx(
        DOCS_DIR / "SaralPrivacy_Gaps_and_TestingPlan.md",
        DOCS_DIR / "SaralPrivacy_Gaps_and_TestingPlan.docx",
        "Gap Analysis & Testing Plan",
    )
    print("Done.")
